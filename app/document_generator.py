"""
Word-document generatie voor beleidsstukken per standaard.

Genereert een .docx met titelpagina en per-eis secties met PDCA-velden.
"""

import io
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


MAANDEN_NL = [
    "", "januari", "februari", "maart", "april", "mei", "juni",
    "juli", "augustus", "september", "oktober", "november", "december",
]

# Kleuren
BLAUW = RGBColor(0x45, 0x99, 0xD5)
DONKER = RGBColor(0x14, 0x19, 0x4A)
GRIJS = RGBColor(0x49, 0x65, 0x80)

PDCA_VELDEN = [
    ("Ambitie", "ambitie"),
    ("Beoogd resultaat", "beoogd_resultaat"),
    ("Concrete acties", "concrete_acties"),
    ("Wijze van meten", "wijze_van_meten"),
]


def generate_beleidsstuk(
    standaard_naam: str,
    eis_lijst: list,
    school_naam: str = "",
    invullingen: dict = None,
) -> bytes:
    """
    Genereer een Word-beleidsstuk voor een standaard.

    Args:
        standaard_naam: Naam van de standaard (bijv. "OP4 - Onderwijstijd")
        eis_lijst: Lijst van (eis_id, eis_data) tuples
        school_naam: Naam van de school
        invullingen: Dict van eis_id -> dict met PDCA-velden (optioneel)

    Returns:
        Bytes van het gegenereerde .docx bestand
    """
    doc = Document()

    # Standaard margins instellen
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _add_title_page(doc, standaard_naam, school_naam)

    for eis_id, eis_data in eis_lijst:
        eis_invulling = (invullingen or {}).get(eis_id, {})
        _add_eis_section(doc, eis_id, eis_data, eis_invulling)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_title_page(doc: Document, standaard_naam: str, school_naam: str):
    """Voeg titelpagina toe aan het document."""
    # Witruimte boven
    for _ in range(6):
        doc.add_paragraph()

    # "Beleidsstuk" titel
    p_titel = doc.add_paragraph()
    p_titel.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_titel.add_run("Beleidsstuk")
    run.font.size = Pt(28)
    run.font.color.rgb = BLAUW
    run.font.bold = True

    # Standaardnaam
    p_standaard = doc.add_paragraph()
    p_standaard.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_standaard.add_run(standaard_naam)
    run.font.size = Pt(22)
    run.font.color.rgb = DONKER
    run.font.bold = True

    # Scheidingslijn
    p_lijn = doc.add_paragraph()
    p_lijn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_lijn.add_run("_" * 50)
    run.font.color.rgb = GRIJS
    run.font.size = Pt(10)

    # Witruimte
    doc.add_paragraph()

    # Schoolnaam
    p_school = doc.add_paragraph()
    p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if school_naam.strip():
        run = p_school.add_run(school_naam)
        run.font.size = Pt(16)
        run.font.color.rgb = GRIJS
    else:
        run = p_school.add_run("[Schoolnaam niet ingevuld]")
        run.font.size = Pt(16)
        run.font.color.rgb = GRIJS
        run.italic = True

    # Datum
    vandaag = date.today()
    datum_str = f"{vandaag.day} {MAANDEN_NL[vandaag.month]} {vandaag.year}"
    p_datum = doc.add_paragraph()
    p_datum.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_datum.add_run(datum_str)
    run.font.size = Pt(12)
    run.font.color.rgb = GRIJS

    # Page break
    doc.add_page_break()


def _add_eis_section(doc: Document, eis_id: str, eis_data: dict, invulling: dict):
    """Voeg een sectie toe voor één eis."""
    titel = eis_data.get("titel", "")
    heading = doc.add_heading(f"{eis_id} - {titel}", level=2)
    for run in heading.runs:
        run.font.color.rgb = DONKER

    # Eisomschrijving
    eisomschrijving = eis_data.get("eisomschrijving", "")
    if eisomschrijving:
        p_label = doc.add_paragraph()
        run = p_label.add_run("Eisomschrijving")
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = GRIJS

        p_omschr = doc.add_paragraph(eisomschrijving)
        p_omschr.style.font.size = Pt(10)

    # PDCA-velden
    for label, key in PDCA_VELDEN:
        doc.add_paragraph()  # witruimte

        p_label = doc.add_paragraph()
        run = p_label.add_run(label)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = BLAUW

        waarde = invulling.get(key, "")
        if waarde and waarde.strip():
            doc.add_paragraph(waarde)
        else:
            p_placeholder = doc.add_paragraph()
            run = p_placeholder.add_run("[Nog niet ingevuld]")
            run.italic = True
            run.font.color.rgb = GRIJS

    # Ruimte na eis-sectie
    doc.add_paragraph()


# =============================================================================
# AI-gegenereerd beleidsstuk
# =============================================================================

def generate_ai_beleidsstuk(
    standaard_naam: str,
    school_naam: str,
    hoofdstukken: list,
    metadata: dict = None,
) -> bytes:
    """
    Genereer een Word-beleidsstuk met AI-gegenereerde hoofdstukken.

    Args:
        standaard_naam: Naam van de standaard
        school_naam: Naam van de school
        hoofdstukken: List[HoofdstukResultaat] met gegenereerde content
        metadata: Optioneel dict met versie, verantwoordelijke, etc.

    Returns:
        Bytes van het gegenereerde .docx bestand
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _add_title_page(doc, standaard_naam, school_naam)
    _add_metadata_table(doc, metadata)
    _add_table_of_contents(doc, hoofdstukken)

    nummer = 1
    for hoofdstuk in hoofdstukken:
        if hoofdstuk.skipped and not hoofdstuk.content:
            continue
        _add_ai_chapter(doc, nummer, hoofdstuk)
        nummer += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_metadata_table(doc: Document, metadata: dict = None):
    """Voeg metadata-tabel toe (versie, datum, verantwoordelijke, etc.)."""
    meta = metadata or {}

    vandaag = date.today()
    datum_str = f"{vandaag.day} {MAANDEN_NL[vandaag.month]} {vandaag.year}"

    rows_data = [
        ("Versie", meta.get("versie", "[In te vullen]")),
        ("Datum", datum_str),
        ("Verantwoordelijke", meta.get("verantwoordelijke", "[In te vullen]")),
        ("Vaststelling", meta.get("vaststelling", "[In te vullen]")),
        ("Verloopdatum", meta.get("verloopdatum", "[In te vullen]")),
        ("MR", meta.get("mr", "[In te vullen]")),
    ]

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Light Grid Accent 1"

    for i, (label, waarde) in enumerate(rows_data):
        cell_label = table.rows[i].cells[0]
        cell_waarde = table.rows[i].cells[1]

        cell_label.text = ""
        run = cell_label.paragraphs[0].add_run(label)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = DONKER

        cell_waarde.text = ""
        run = cell_waarde.paragraphs[0].add_run(waarde)
        run.font.size = Pt(10)
        if waarde == "[In te vullen]":
            run.font.color.rgb = GRIJS
            run.italic = True

    doc.add_paragraph()
    doc.add_page_break()


def _add_table_of_contents(doc: Document, hoofdstukken: list):
    """Voeg handmatige inhoudsopgave toe."""
    heading = doc.add_heading("Inhoudsopgave", level=1)
    for run in heading.runs:
        run.font.color.rgb = DONKER

    doc.add_paragraph()

    nummer = 1
    for hoofdstuk in hoofdstukken:
        if hoofdstuk.skipped and not hoofdstuk.content:
            continue
        p = doc.add_paragraph()
        run = p.add_run(f"{nummer}. {hoofdstuk.label}")
        run.font.size = Pt(12)
        run.font.color.rgb = DONKER
        nummer += 1

    doc.add_page_break()


def _strip_duplicate_heading(content: str, label: str) -> str:
    """Strip de hoofdstuktitel als die door de AI als eerste regel is herhaald."""
    lines = content.split("\n")
    if not lines:
        return content
    first = lines[0].strip()
    # Check diverse vormen: "Ambitie", "## Ambitie", "**Ambitie**", "# Ambitie"
    cleaned = first.lstrip("#").strip().strip("*").strip()
    if cleaned.lower() == label.lower():
        return "\n".join(lines[1:]).lstrip("\n")
    return content


def _add_ai_chapter(doc: Document, nummer: int, hoofdstuk):
    """Voeg een AI-gegenereerd hoofdstuk toe."""
    # Heading
    heading = doc.add_heading(f"{nummer}. {hoofdstuk.label}", level=1)
    for run in heading.runs:
        run.font.color.rgb = DONKER

    content = hoofdstuk.content

    # Strip duplicate heading als de AI de hoofdstuktitel herhaalt
    content = _strip_duplicate_heading(content, hoofdstuk.label)

    # Placeholder tekst
    if hoofdstuk.skipped or not content or content == "[Nog niet ingevuld door de school]":
        p = doc.add_paragraph()
        run = p.add_run("[Nog niet ingevuld door de school]")
        run.italic = True
        run.font.color.rgb = GRIJS
        doc.add_paragraph()
        return

    # Error tekst
    if hoofdstuk.error:
        p = doc.add_paragraph()
        run = p.add_run(f"[Fout bij genereren: {hoofdstuk.error}]")
        run.italic = True
        run.font.color.rgb = GRIJS
        doc.add_paragraph()
        return

    # Parse AI output: alinea's en bullet points
    _add_ai_content(doc, content)

    doc.add_paragraph()


def _add_ai_content(doc: Document, content: str):
    """Parse AI-gegenereerde tekst en voeg toe met juiste formatting."""
    lines = content.split("\n")
    current_paragraph_lines = []

    def flush_paragraph():
        """Schrijf opgespaarde niet-bullet regels als alinea."""
        if current_paragraph_lines:
            text = " ".join(current_paragraph_lines)
            _add_body_paragraph(doc, text)
            current_paragraph_lines.clear()

    for line in lines:
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            continue

        # Markdown headings (## of ###) -> sub-heading
        if stripped.startswith("## ") or stripped.startswith("### "):
            flush_paragraph()
            heading_text = stripped.lstrip("#").strip()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = BLAUW

        # Bullet point
        elif stripped.startswith("- ") or stripped.startswith("* "):
            flush_paragraph()
            bullet_text = stripped[2:]
            _add_bullet_paragraph(doc, bullet_text)

        # Genummerde lijst (bijv. "1. ", "2. ")
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == "." and stripped[2] == " ":
            flush_paragraph()
            _add_bullet_paragraph(doc, stripped)

        # Sub-heading (lijnen die beginnen met ** en eindigen met **)
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            flush_paragraph()
            heading_text = stripped[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = BLAUW

        else:
            # Verwijder markdown bold markers uit lopende tekst
            clean = stripped.replace("**", "")
            current_paragraph_lines.append(clean)

    flush_paragraph()


def _add_body_paragraph(doc: Document, text: str):
    """Voeg een body-alinea toe."""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)


def _add_bullet_paragraph(doc: Document, text: str):
    """Voeg een bullet-point alinea toe."""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
